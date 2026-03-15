"""Document parsers — returns list[Page] for all supported file types."""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from core.logger import logger


@dataclass
class Page:
    """Single page / section of a parsed document."""
    page_number: int
    text: str
    has_tables: bool = False
    tables: list[str] = field(default_factory=list)   # raw table text representations
    metadata: dict = field(default_factory=dict)


# ── PDF ───────────────────────────────────────────────────────────────────────

def parse_pdf(file_path: str) -> list[Page]:
    """
    Parse a PDF using PyMuPDF for text and pdfplumber for table detection.
    Returns one Page per PDF page.
    """
    import fitz  # PyMuPDF
    import pdfplumber

    pages: list[Page] = []

    # Open with pdfplumber for table extraction
    with pdfplumber.open(file_path) as pdf_pb:
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            fitz_page = doc[page_num]
            text = fitz_page.get_text("text") or ""

            tables: list[str] = []
            has_tables = False

            # Table detection via pdfplumber
            try:
                pb_page = pdf_pb.pages[page_num]
                extracted_tables = pb_page.extract_tables()
                if extracted_tables:
                    has_tables = True
                    for table in extracted_tables:
                        # Render table as plain text (CSV-like)
                        table_lines = []
                        for row in table:
                            cells = [str(cell or "").strip() for cell in row]
                            table_lines.append(" | ".join(cells))
                        tables.append("\n".join(table_lines))
            except Exception as exc:
                logger.debug(f"Table extraction failed on page {page_num + 1}: {exc}")

            pages.append(
                Page(
                    page_number=page_num + 1,
                    text=text.strip(),
                    has_tables=has_tables,
                    tables=tables,
                    metadata={"source_type": "pdf", "total_pages": len(doc)},
                )
            )

        doc.close()

    return pages


# ── DOCX ──────────────────────────────────────────────────────────────────────

def parse_docx(file_path: str) -> list[Page]:
    """
    Parse a DOCX file using python-docx.
    Splits into virtual pages of ~3000 characters to mirror PDF paging.
    """
    from docx import Document

    CHARS_PER_PAGE = 3000
    doc = Document(file_path)

    full_text_parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    # Include table text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            full_text_parts.append(" | ".join(cells))

    full_text = "\n".join(full_text_parts)

    if not full_text.strip():
        return [Page(page_number=1, text="(empty document)", metadata={"source_type": "docx"})]

    pages: list[Page] = []
    for i in range(0, len(full_text), CHARS_PER_PAGE):
        chunk = full_text[i : i + CHARS_PER_PAGE]
        page_num = i // CHARS_PER_PAGE + 1
        pages.append(
            Page(
                page_number=page_num,
                text=chunk.strip(),
                has_tables=len(doc.tables) > 0,
                metadata={"source_type": "docx"},
            )
        )

    return pages


# ── TXT ───────────────────────────────────────────────────────────────────────

def parse_txt(file_path: str) -> list[Page]:
    """
    Parse a plain text file.
    Splits into virtual pages of ~3000 characters.
    """
    CHARS_PER_PAGE = 3000

    with open(file_path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    if not content.strip():
        return [Page(page_number=1, text="(empty file)", metadata={"source_type": "txt"})]

    pages: list[Page] = []
    for i in range(0, len(content), CHARS_PER_PAGE):
        chunk = content[i : i + CHARS_PER_PAGE]
        page_num = i // CHARS_PER_PAGE + 1
        pages.append(
            Page(
                page_number=page_num,
                text=chunk.strip(),
                metadata={"source_type": "txt"},
            )
        )

    return pages


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(file_path: str, mime_type: str | None = None) -> list[Page]:
    """Route to the correct parser based on file extension / MIME type."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf" or (mime_type and "pdf" in mime_type):
        return parse_pdf(file_path)
    if ext in (".docx", ".doc") or (mime_type and "word" in mime_type):
        return parse_docx(file_path)
    if ext == ".txt" or (mime_type and "text/plain" in mime_type):
        return parse_txt(file_path)

    # Fallback to plain text
    logger.warning(f"Unknown file type '{ext}', falling back to txt parser")
    return parse_txt(file_path)
